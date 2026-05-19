from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).parent
SOURCE = ROOT / "ASTRAEL WIKI.md"
MAP_IMAGE = ROOT / "aresia-continent-map-v2.png"
OUTPUT = ROOT / "ASTRAEL WIKI.docx"

ACCENT = "17324D"
ACCENT_2 = "C6A15B"
LIGHT = "EEF3F7"
QUOTE = "F6F1E8"
TEXT = "24313C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 18, ACCENT),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 11.5, ACCENT),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = "Arial"
    list_bullet.font.size = Pt(10.5)
    list_bullet.paragraph_format.space_after = Pt(2)

    list_number = doc.styles["List Number"]
    list_number.font.name = "Arial"
    list_number.font.size = Pt(10.5)
    list_number.paragraph_format.space_after = Pt(2)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("ASTRAEL WIKI")
    set_run_font(run, size=28, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ภาค 1 - The Fall and Rebirth")
    set_run_font(run, size=14, bold=True, color=ACCENT_2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("คู่มือโลก ตัวละคร อาณาจักร และระบบพลัง")
    set_run_font(run, size=11, color=TEXT)

    doc.add_paragraph()
    doc.add_picture(str(MAP_IMAGE), width=Inches(6.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("แผนที่ทวีปอาเรเซีย")
    set_run_font(run, size=9.5, color="5B6770")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("อ้างอิงจากเนื้อหาตอนที่ 1-42")
    set_run_font(run, size=9.5, color="5B6770")

    doc.add_page_break()


def extract_toc_entries(lines: list[str]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for line in lines:
        match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2).strip()
        if text == "ASTRAEL WIKI":
            continue
        if level <= 2:
            entries.append((level, text))
    return entries


def add_toc(doc: Document, entries: list[tuple[int, str]]) -> None:
    h = doc.add_heading("สารบัญ", level=1)
    h.paragraph_format.space_after = Pt(8)
    for level, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.0 if level == 1 else 0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=10.5 if level == 1 else 10, bold=level == 1, color=ACCENT if level == 1 else TEXT)
    doc.add_page_break()


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), QUOTE)
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=ACCENT)
    run.italic = True


def add_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_2)
    pbdr.append(bottom)
    pPr.append(pbdr)


def parse_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=ACCENT)
            run.bold = True
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set()


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, cell_text in enumerate(rows[0]):
        cell = header_cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell)
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": ACCENT})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        set_run_font(run, size=9.5, bold=True, color="FFFFFF")

    for row_index, row in enumerate(rows[1:], start=1):
        cells = table.add_row().cells
        fill = "FFFFFF" if row_index % 2 else LIGHT
        for i, cell_text in enumerate(row):
            cell = cells[i]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "D7DEE5"},
                bottom={"val": "single", "sz": "4", "color": "D7DEE5"},
                left={"val": "single", "sz": "4", "color": "D7DEE5"},
                right={"val": "single", "sz": "4", "color": "D7DEE5"},
            )
            p = cell.paragraphs[0]
            parse_inline(p, cell_text)
    doc.add_paragraph()


def add_body(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            add_rule(doc)
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if text != "ASTRAEL WIKI":
                doc.add_heading(text, level=min(level, 3))
            i += 1
            continue
        if stripped.startswith("> "):
            add_quote(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [split_table_row(row) for row in table_lines if not is_table_separator(row)]
            add_markdown_table(doc, rows)
            continue
        bullet = re.match(r"^\*\s+(.*)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, numbered.group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        parse_inline(p, stripped.replace("  ", " "))
        i += 1


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_toc(doc, extract_toc_entries(lines))
    add_body(doc, lines)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
